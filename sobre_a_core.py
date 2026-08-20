"""
sobre_a_core.py — Fase 4 del Expediente Digital v2.

Lógica pura de generación del Sobre A:
- Catálogo de piezas (formularios SNCC generables + documentos del expediente).
- Estado por pieza aplicando resolver_estado() (regla D2): documento vencido
  JAMÁS entra al ZIP; pendiente entra marcado; rechazado/faltante bloqueado.
- Generadores docx con python-docx: SIN logo institucional, con el nombre de
  la institución del proceso parametrizado en la cabecera.
- Huecos con sombreado [[PENDIENTE: campo]] cuando falta un dato.
- Firma: si el firmante principal (bid_representantes.es_firmante_principal)
  tiene firma_url, se inserta la imagen sobre la línea de firma; si no, la
  línea queda en blanco para firma manuscrita.
- Hojas separadoras (reportlab) antes de cada documento adjunto; si el
  documento falta, queda SOLO el separador con la marca de pendiente.
- Ensamblado del ZIP y, si LibreOffice está disponible y formato='pdf',
  conversión de formularios a PDF + merge completo con carátula/índice
  paginado contando páginas reales (pypdf) DESPUÉS del merge.

Este módulo NO conoce FastAPI: recibe el cliente de Supabase y devuelve
bytes/estructuras. El router (router_sobre_a.py) orquesta jobs y auth.
"""

from __future__ import annotations

import io
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from datetime import date, datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import cm
from reportlab.pdfgen import canvas as rl_canvas

from expediente_core import resolver_estado

# ══════════════════════════════════════════════════════════════
# Utilidades de datos
# ══════════════════════════════════════════════════════════════

MARCA_PENDIENTE = "[[PENDIENTE: {campo}]]"


def _hoy() -> str:
    return date.today().strftime("%d/%m/%Y")


def _v(registro: dict | None, *campos: str) -> str | None:
    """Primer campo no vacío entre equivalentes (patrón del resumen D1)."""
    if not registro:
        return None
    for c in campos:
        val = registro.get(c)
        if val is not None and str(val).strip():
            return str(val).strip()
    return None


def _vencido(fecha_vencimiento) -> bool:
    if not fecha_vencimiento:
        return False
    try:
        fv = str(fecha_vencimiento)[:10]
        return date.fromisoformat(fv) < date.today()
    except (ValueError, TypeError):
        return False


def _slug(texto: str, maximo: int = 60) -> str:
    limpio = re.sub(r"[^A-Za-z0-9._-]+", "_", texto or "").strip("_")
    return (limpio or "documento")[:maximo]


# ══════════════════════════════════════════════════════════════
# Contexto: datos maestros para llenar formularios
# ══════════════════════════════════════════════════════════════

def datos_contexto(sb, empresa_id: str, referencia: str) -> dict:
    """
    Junta en un dict todo lo necesario para autollenar:
    empresa (perfiles_empresa), proceso (bid_matching_pliegos),
    firmante principal (bid_representantes) y bytes de su firma si existe.
    """
    emp = (sb.table("perfiles_empresa").select("*")
           .eq("id", empresa_id).limit(1).execute().data or [None])[0] or {}

    proc = (sb.table("bid_matching_pliegos")
            .select("referencia,nombre_proceso,institucion,presupuesto_base")
            .eq("empresa_id", empresa_id).eq("referencia", referencia)
            .order("creado_en", desc=True).limit(1).execute().data or [None])[0] or {}

    firmante = (sb.table("bid_representantes").select("*")
                .eq("empresa_id", empresa_id).eq("activo", True)
                .eq("es_firmante_principal", True)
                .limit(1).execute().data or [None])[0]
    if not firmante:  # sin principal marcado: el primero activo
        firmante = (sb.table("bid_representantes").select("*")
                    .eq("empresa_id", empresa_id).eq("activo", True)
                    .order("creado_en").limit(1).execute().data or [None])[0] or {}

    firma_png = None
    if firmante.get("firma_url"):
        firma_png = _descargar_binario(sb, firmante["firma_url"])

    return {
        "empresa": emp,
        "proceso": proc,
        "firmante": firmante,
        "firma_png": firma_png,
        "referencia": referencia,
        "institucion": _v(proc, "institucion") or None,
    }


def _descargar_binario(sb, ruta: str) -> bytes | None:
    """Descarga un archivo del bucket (ruta interna) o de una URL http."""
    try:
        if ruta.startswith("http"):
            import httpx
            r = httpx.get(ruta, timeout=20, follow_redirects=True)
            return r.content if r.status_code == 200 else None
        from router_bid_manager import STORAGE_BUCKET
        return sb.storage.from_(STORAGE_BUCKET).download(ruta)
    except Exception:
        return None


# ══════════════════════════════════════════════════════════════
# Catálogo de piezas
# ══════════════════════════════════════════════════════════════
# Formularios generables. Cada clave: (código SNCC, título, fn generadora)
# — la fn se resuelve al final del módulo para poder definirlas abajo.

FORMULARIOS = {
    "f034": ("SNCC.F.034", "Formulario de Presentación de Oferta"),
    "f042": ("SNCC.F.042", "Formulario de Información sobre el Oferente"),
    "f035": ("SNCC.F.035", "Estructura para brindar soporte técnico al Equipo ofertado"),
    "f036": ("SNCC.F.036", "Listado de Equipos del Oferente"),
    "f037": ("SNCC.F.037", "Personal de Plantilla del Oferente"),
    "d045": ("SNCC.D.045", "Currículo del Personal Profesional propuesto"),
    "d049": ("SNCC.D.049", "Experiencia como Contratista"),
    "djur": ("Declaración Jurada", "Declaración Jurada simple (Art. 14 Ley 340-06 / Art. 38 Ley 47-25)"),
}

# Documentos del expediente que suelen exigir los pliegos (visto en los
# análisis de bid_matching_pliegos). Cada entrada define de dónde salen.
#   entidad -> secciones de bid_archivos / campos legacy
DOCS_EXPEDIENTE = [
    ("certificaciones", "Certificaciones y registros",
     "Certificaciones vigentes (RPE, DGII, TSS, Registro Mercantil, CODIA, MIPYME, no antecedentes…)"),
    ("financieros", "Estados financieros",
     "Estados financieros auditados de los últimos ejercicios"),
    ("representantes", "Representantes legales",
     "Cédulas y poderes de representación"),
    ("socios", "Socios y accionistas",
     "Cédulas, nómina de accionistas y acta de asamblea"),
]


def listar_piezas(sb, empresa_id: str, referencia: str) -> list[dict]:
    """
    Estado por pieza para el modal de generación.
    Devuelve: [{id, tipo, nombre, estado, motivo, seccion}]
      - tipo 'formulario': siempre generable; estado warn si le faltan datos
        clave (se generará con huecos), ok si el contexto está completo.
      - tipo 'documento': un ítem por adjunto validable; estado por D2.
        Vencido => err con motivo explícito (JAMÁS entra al ZIP).
    """
    ctx = datos_contexto(sb, empresa_id, referencia)
    piezas: list[dict] = []

    # --- Formularios ---
    for clave, (codigo, titulo) in FORMULARIOS.items():
        faltan = _campos_faltantes(sb, empresa_id, clave, ctx)
        piezas.append({
            "id": clave,
            "tipo": "formulario",
            "nombre": f"{codigo} — {titulo}",
            "estado": "warn" if faltan else "ok",
            "motivo": (f"Se generará con huecos: {', '.join(faltan[:4])}"
                       + ("…" if len(faltan) > 4 else "")) if faltan else None,
            "seccion": "Formularios",
        })

    # --- Documentos del expediente ---
    for entidad, seccion, _desc in DOCS_EXPEDIENTE:
        for a in _adjuntos_entidad(sb, empresa_id, entidad):
            piezas.append(a | {"seccion": seccion})

    return piezas


def _campos_faltantes(sb, empresa_id: str, clave: str, ctx: dict) -> list[str]:
    """Qué datos le faltarían a un formulario (para avisar antes de generar)."""
    emp, fir = ctx["empresa"], ctx["firmante"]
    faltan = []

    def req(valor, etiqueta):
        if not valor:
            faltan.append(etiqueta)

    if clave in ("f034", "f042", "djur"):
        req(_v(emp, "razon_social", "nombre_perfil"), "Razón social")
        req(_v(emp, "rnc"), "RNC")
        req(_v(fir, "nombre_completo"), "Representante")
    if clave == "f042":
        req(_v(emp, "rpe"), "RPE")
        req(_v(emp, "direccion_completa", "domicilio", "direccion"), "Dirección")
        req(_v(emp, "email_empresa"), "Correo institucional")
    if clave == "f036":
        n = (sb.table("bid_equipos").select("id", count="exact")
             .eq("empresa_id", empresa_id).eq("activo", True).execute().count or 0)
        if n == 0:
            faltan.append("Equipos registrados")
    if clave in ("f037", "d045"):
        n = (sb.table("bid_personal").select("id", count="exact")
             .eq("empresa_id", empresa_id).eq("activo", True).execute().count or 0)
        if n == 0:
            faltan.append("Personal registrado")
    if clave == "d049":
        n = (sb.table("bid_experiencia").select("id", count="exact")
             .eq("empresa_id", empresa_id).execute().count or 0)
        if n == 0:
            faltan.append("Experiencia registrada")
    if clave == "f035":
        faltan.append("Descripción de la estructura de soporte")
    return faltan


def _adjuntos_entidad(sb, empresa_id: str, entidad: str) -> list[dict]:
    """Piezas 'documento' de una entidad, con estado D2 por adjunto."""
    from expediente_core import TABLA_POR_ENTIDAD
    tabla = TABLA_POR_ENTIDAD[entidad]
    registros = (sb.table(tabla).select("*")
                 .eq("empresa_id", empresa_id).execute().data or [])
    salida = []
    for r in registros:
        vencido = _vencido(r.get("fecha_vencimiento"))
        etiqueta = (_v(r, "nombre_display", "tipo") or _v(r, "periodo")
                    or _v(r, "nombre_completo") or "Documento")

        adjuntos = (sb.table("bid_archivos").select("id,nombre,estado,eliminado_en")
                    .eq("empresa_id", empresa_id).eq("entidad", entidad)
                    .eq("registro_id", r["id"]).is_("eliminado_en", "null")
                    .execute().data or [])
        # Legacy: certificaciones con archivo_url directo
        if not adjuntos and entidad == "certificaciones" and r.get("archivo_url"):
            adjuntos = [{"id": f"legacy:{r['id']}", "nombre": r.get("archivo_nombre") or etiqueta,
                         "estado": "pendiente", "_legacy_path": r["archivo_url"]}]

        if not adjuntos:
            salida.append({
                "id": f"doc:{entidad}:{r['id']}:falta",
                "tipo": "documento", "nombre": etiqueta,
                "estado": "err",
                "motivo": "Sin archivo adjunto — entrará solo la hoja separadora",
                "registro_id": r["id"], "entidad": entidad,
            })
            continue

        for a in adjuntos:
            est = resolver_estado(a.get("estado"), vencido=vencido)
            motivo = None
            if vencido:
                motivo = "Documento vencido — bloqueado, no entra al ZIP"
            elif est == "warn":
                motivo = "Sin verificar — entra marcado como pendiente de validación"
            elif est == "err":
                motivo = "Rechazado — bloqueado"
            salida.append({
                "id": f"doc:{entidad}:{r['id']}:{a['id']}",
                "tipo": "documento",
                "nombre": f"{etiqueta} — {a.get('nombre') or 'archivo'}",
                "estado": est, "motivo": motivo,
                "registro_id": r["id"], "entidad": entidad,
                "archivo_id": a["id"],
                "_legacy_path": a.get("_legacy_path"),
            })
    return salida


# ══════════════════════════════════════════════════════════════
# Bloques docx reutilizables (sin logo, institución parametrizada)
# ══════════════════════════════════════════════════════════════

def _sombrear(parrafo, color="FFF2CC"):
    """Sombreado amarillo suave para los huecos [[PENDIENTE]]."""
    for run in parrafo.runs:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), color)
        run._element.get_or_add_rPr().append(shd)


def _p(doc, texto="", *, bold=False, size=11, center=False, upper=False):
    par = doc.add_paragraph()
    run = par.add_run(texto.upper() if upper else texto)
    run.bold = bold
    run.font.size = Pt(size)
    if center:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "[[PENDIENTE" in texto:
        _sombrear(par)
    return par


def _dato(ctx_reg, etiqueta: str, *campos) -> str:
    """Valor o marca de pendiente sombreable."""
    return _v(ctx_reg, *campos) or MARCA_PENDIENTE.format(campo=etiqueta)


def _cabecera(doc, ctx: dict, codigo: str, titulo: str):
    """
    Cabecera estándar SIN logo: código SNCC, institución del proceso
    (parametrizada — regla del cliente), referencia y fecha.
    """
    inst = ctx.get("institucion") or MARCA_PENDIENTE.format(campo="Institución del proceso")
    _p(doc, codigo, bold=True, size=10)
    _p(doc, inst, bold=True, size=13, center=True, upper=True)
    _p(doc, titulo, bold=True, size=12, center=True, upper=True)
    _p(doc, f"No. EXPEDIENTE: {ctx['referencia']}", size=10, center=True)
    _p(doc, f"Fecha: {_hoy()}", size=10, center=True)
    doc.add_paragraph()


def _bloque_firma(doc, ctx: dict):
    """Línea de firma; inserta imagen de firma si el firmante la tiene."""
    fir = ctx["firmante"]
    doc.add_paragraph()
    nombre = _dato(fir, "Nombre del representante", "nombre_completo")
    cargo = _v(fir, "cargo") or "Representante Legal"
    razon = _dato(ctx["empresa"], "Razón social", "razon_social", "nombre_perfil")
    _p(doc, f"(Nombre y apellido) {nombre}, en calidad de {cargo}, "
            f"debidamente autorizado para actuar en nombre y representación de {razon}.")
    doc.add_paragraph()

    if ctx.get("firma_png"):
        try:
            par = doc.add_paragraph()
            par.add_run().add_picture(io.BytesIO(ctx["firma_png"]), width=Inches(1.8))
        except Exception:
            pass  # imagen corrupta: se queda la línea en blanco
    _p(doc, "Firma ____________________________________")
    _p(doc, "Sello")


def _tabla(doc, encabezados: list[str], filas: list[list[str]]):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Table Grid"
    for i, h in enumerate(encabezados):
        cel = t.rows[0].cells[i]
        cel.text = h
        for run in cel.paragraphs[0].runs:
            run.bold = True
    for fila in filas:
        celdas = t.add_row().cells
        for i, val in enumerate(fila):
            celdas[i].text = str(val) if val is not None else ""
    return t


def _docx_bytes(doc) -> bytes:
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════
# Generadores de formularios
# ══════════════════════════════════════════════════════════════

def generar_f034(sb, empresa_id: str, ctx: dict) -> bytes:
    doc = Document()
    _cabecera(doc, ctx, "SNCC.F.034", "Presentación de Oferta")
    inst = ctx.get("institucion") or MARCA_PENDIENTE.format(campo="Institución")
    proc = ctx["proceso"]
    objeto = _v(proc, "nombre_proceso") or MARCA_PENDIENTE.format(campo="Objeto del proceso")

    _p(doc, "Señores")
    _p(doc, inst, bold=True, upper=True)
    doc.add_paragraph()
    _p(doc, "Nosotros, los suscritos, declaramos que:")
    for texto in [
        "Hemos examinado y no tenemos reservas al Pliego de Condiciones del "
        f"procedimiento de referencia {ctx['referencia']}, incluyendo sus "
        "enmiendas/adendas: ____________________________________________",
        "De conformidad con el Pliego de Condiciones y el Cronograma de "
        f"Ejecución, nos comprometemos a ejecutar: {objeto}.",
        "Si nuestra oferta es aceptada, nos comprometemos a obtener una garantía "
        "de fiel cumplimiento del Contrato, de conformidad con el Pliego de "
        "Condiciones, por el importe requerido para asegurar su fiel cumplimiento.",
        "Para este procedimiento no somos partícipes en calidad de Oferentes en "
        "más de una Oferta, de conformidad con el Pliego de Condiciones.",
        "Nuestra firma, sus afiliadas o subsidiarias, incluyendo cualquier "
        "subcontratista o proveedor, no han sido declarados inelegibles para "
        "presentar ofertas.",
        "Entendemos que esta Oferta, junto con su aceptación por escrito incluida "
        "en la notificación de adjudicación, constituirá una obligación "
        "contractual hasta la preparación y ejecución del Contrato formal.",
        "Entendemos que la Entidad Contratante no está obligada a aceptar la "
        "Oferta evaluada como la más baja ni ninguna otra de las que reciba.",
    ]:
        doc.add_paragraph(texto, style="List Bullet")
    _bloque_firma(doc, ctx)
    return _docx_bytes(doc)


def generar_f042(sb, empresa_id: str, ctx: dict) -> bytes:
    doc = Document()
    _cabecera(doc, ctx, "SNCC.F.042", "Formulario de Información sobre el Oferente")
    emp, fir = ctx["empresa"], ctx["firmante"]
    filas = [
        ["1. Nombre / Razón Social del Oferente",
         _dato(emp, "Razón social", "razon_social", "nombre_perfil")],
        ["2. Si es consorcio, nombre jurídico de cada miembro", "N/A"],
        ["3. RNC / Cédula del Oferente", _dato(emp, "RNC", "rnc")],
        ["4. RPE del Oferente", _dato(emp, "RPE", "rpe")],
        ["5. Domicilio legal del Oferente",
         _dato(emp, "Dirección", "direccion_completa", "domicilio", "direccion")],
        ["6. Representante autorizado",
         f"Nombre: {_dato(fir, 'Nombre', 'nombre_completo')}\n"
         f"Cédula: {_dato(fir, 'Cédula', 'cedula')}\n"
         f"Dirección: {_dato(fir, 'Dirección', 'direccion')}\n"
         f"Teléfono: {_dato(fir, 'Teléfono', 'telefono')}\n"
         f"Correo electrónico: {_dato(fir, 'Correo', 'email')}"],
        ["7. Correo institucional", _dato(emp, "Correo institucional", "email_empresa")],
        ["8. Teléfono", _dato(emp, "Teléfono", "telefono_principal")],
    ]
    _tabla(doc, ["Campo", "Información"], filas)
    _bloque_firma(doc, ctx)
    return _docx_bytes(doc)


def generar_f036(sb, empresa_id: str, ctx: dict) -> bytes:
    doc = Document()
    _cabecera(doc, ctx, "SNCC.F.036", "Listado de Equipos del Oferente")
    equipos = (sb.table("bid_equipos").select("*")
               .eq("empresa_id", empresa_id).eq("activo", True)
               .order("descripcion").execute().data or [])
    if not equipos:
        _p(doc, MARCA_PENDIENTE.format(campo="Equipos registrados en el expediente"))
    else:
        filas = [[
            e.get("cantidad") or 1,
            _v(e, "descripcion") or "",
            " ".join(x for x in [_v(e, "marca"), _v(e, "modelo"), _v(e, "anio")] if x),
            _v(e, "capacidad") or "",
            (_v(e, "propiedad") or "").capitalize(),
            _v(e, "empresa_alquiler") or "",
        ] for e in equipos]
        _tabla(doc, ["Cant.", "Descripción", "Marca / Modelo / Año",
                     "Capacidad", "Propiedad", "Empresa de alquiler"], filas)
        _p(doc, "\nNota: se anexan certificados de propiedad, contratos de alquiler "
                "y/o declaraciones juradas de disponibilidad, según aplique.", size=9)
    _bloque_firma(doc, ctx)
    return _docx_bytes(doc)


def generar_f037(sb, empresa_id: str, ctx: dict) -> bytes:
    doc = Document()
    _cabecera(doc, ctx, "SNCC.F.037", "Personal de Plantilla del Oferente")
    personal = (sb.table("bid_personal").select("*")
                .eq("empresa_id", empresa_id).eq("activo", True)
                .order("nombre_completo").execute().data or [])
    if not personal:
        _p(doc, MARCA_PENDIENTE.format(campo="Personal registrado en el expediente"))
    else:
        filas = [[
            _v(p, "nombre_completo") or "",
            _v(p, "cedula") or "",
            _v(p, "cargo_empresa") or "",
            _v(p, "profesion") or "",
            _v(p, "codia") or "",
            p.get("experiencia_general_anios") or "",
        ] for p in personal]
        _tabla(doc, ["Nombre completo", "Cédula", "Cargo", "Profesión",
                     "CODIA", "Años exp."], filas)
    _bloque_firma(doc, ctx)
    return _docx_bytes(doc)


def generar_d045(sb, empresa_id: str, ctx: dict) -> bytes:
    """Currículo del personal profesional: una sección por persona."""
    doc = Document()
    _cabecera(doc, ctx, "SNCC.D.045", "Currículo del Personal Profesional propuesto")
    personal = (sb.table("bid_personal").select("*")
                .eq("empresa_id", empresa_id).eq("activo", True)
                .order("nombre_completo").execute().data or [])
    if not personal:
        _p(doc, MARCA_PENDIENTE.format(campo="Personal registrado en el expediente"))
    for i, p in enumerate(personal):
        if i:
            doc.add_page_break()
        _p(doc, _v(p, "nombre_completo") or "", bold=True, size=12)
        filas = [
            ["Cédula", _dato(p, "Cédula", "cedula")],
            ["Cargo propuesto", _dato(p, "Cargo", "cargo_empresa")],
            ["Profesión", _dato(p, "Profesión", "profesion")],
            ["Formación académica", _dato(p, "Formación", "formacion_academica")],
            ["CODIA", _v(p, "codia") or "N/A"],
            ["Exequátur", _v(p, "exequatur") or "N/A"],
            ["Años de experiencia general", str(p.get("experiencia_general_anios") or
             MARCA_PENDIENTE.format(campo="Experiencia general"))],
            ["Años de experiencia específica", str(p.get("experiencia_especifica_anios") or "N/A")],
            ["Maestría / especialidad",
             (_v(p, "maestria_descripcion") if p.get("tiene_maestria") else "N/A") or "N/A"],
            ["Especialidades", ", ".join(p.get("especialidades") or []) or "N/A"],
            ["Contacto", " / ".join(x for x in [_v(p, "telefono"), _v(p, "email")] if x) or ""],
        ]
        _tabla(doc, ["Campo", "Detalle"], filas)
        _p(doc, "\nSe anexan: CV, cédula, títulos, matrícula CODIA vigente y carta "
                "de intención y disponibilidad, según aplique.", size=9)
    _bloque_firma(doc, ctx)
    return _docx_bytes(doc)


def generar_d049(sb, empresa_id: str, ctx: dict) -> bytes:
    doc = Document()
    _cabecera(doc, ctx, "SNCC.D.049", "Experiencia como Contratista")
    exps = (sb.table("bid_experiencia").select("*")
            .eq("empresa_id", empresa_id)
            .order("fecha_inicio", desc=True).execute().data or [])
    if not exps:
        _p(doc, MARCA_PENDIENTE.format(campo="Experiencia registrada en el expediente"))
    else:
        filas = [[
            _v(e, "nombre_proyecto") or "",
            _v(e, "cliente") or "",
            _v(e, "tipo_obra") or "",
            _v(e, "ubicacion", "provincia") or "",
            f"{_v(e, 'moneda') or 'RD$'} {e.get('monto_contrato') or ''}",
            f"{str(e.get('fecha_inicio') or '')[:10]} — {str(e.get('fecha_fin') or '')[:10]}",
            (_v(e, "estado") or "").capitalize(),
        ] for e in exps]
        _tabla(doc, ["Proyecto", "Contratante", "Tipo de obra", "Ubicación",
                     "Monto", "Inicio — Fin", "Estado"], filas)
        _p(doc, "\nSe anexan certificaciones de experiencia y/o contratos "
                "registrados con entidad contratante, objeto, fechas de inicio "
                "y finalización.", size=9)
    _bloque_firma(doc, ctx)
    return _docx_bytes(doc)


def generar_f035(sb, empresa_id: str, ctx: dict) -> bytes:
    doc = Document()
    _cabecera(doc, ctx, "SNCC.F.035",
              "Estructura para brindar soporte técnico al Equipo ofertado")
    _p(doc, MARCA_PENDIENTE.format(
        campo="Describir la estructura de soporte técnico: talleres, mecánicos, "
              "repuestos, logística de mantenimiento de los equipos ofertados"))
    doc.add_paragraph()
    _p(doc, "_" * 90)
    _p(doc, "_" * 90)
    _p(doc, "_" * 90)
    _bloque_firma(doc, ctx)
    return _docx_bytes(doc)


def generar_djur(sb, empresa_id: str, ctx: dict) -> bytes:
    doc = Document()
    _cabecera(doc, ctx, "DECLARACIÓN JURADA", "Declaración Jurada simple del Oferente")
    emp, fir = ctx["empresa"], ctx["firmante"]
    razon = _dato(emp, "Razón social", "razon_social", "nombre_perfil")
    rnc = _dato(emp, "RNC", "rnc")
    nombre = _dato(fir, "Nombre del representante", "nombre_completo")
    cedula = _dato(fir, "Cédula", "cedula")
    _p(doc,
       f"Quien suscribe, {nombre}, portador(a) de la cédula de identidad y "
       f"electoral No. {cedula}, actuando en nombre y representación de "
       f"{razon}, RNC {rnc}, DECLARA BAJO LA MÁS SOLEMNE FE DEL JURAMENTO que "
       "dicha entidad no se encuentra dentro de las prohibiciones e "
       "inhabilidades establecidas en el artículo 8 numeral 3 y el artículo 14 "
       "de la Ley núm. 340-06 y sus modificaciones, ni en el artículo 38 de la "
       "Ley núm. 47-25 sobre Compras y Contrataciones Públicas, para "
       "participar en procedimientos de contratación con el Estado Dominicano.")
    doc.add_paragraph()
    _p(doc, f"La presente declaración se emite para el procedimiento de "
            f"referencia {ctx['referencia']}, a los {_hoy()}.")
    _bloque_firma(doc, ctx)
    return _docx_bytes(doc)


GENERADORES = {
    "f034": generar_f034, "f042": generar_f042, "f035": generar_f035,
    "f036": generar_f036, "f037": generar_f037, "d045": generar_d045,
    "d049": generar_d049, "djur": generar_djur,
}


# ══════════════════════════════════════════════════════════════
# Separadores, carátula e índice (reportlab)
# ══════════════════════════════════════════════════════════════

def hoja_separadora(nombre_doc: str, referencia: str, institucion: str | None,
                    pendiente: bool = False, motivo: str | None = None) -> bytes:
    """Hoja separadora PDF que precede a cada documento del paquete."""
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    w, h = letter
    c.setLineWidth(2)
    c.rect(1.5 * cm, 1.5 * cm, w - 3 * cm, h - 3 * cm)
    c.setFont("Helvetica", 10)
    c.drawCentredString(w / 2, h - 3 * cm, (institucion or "").upper())
    c.setFont("Helvetica", 11)
    c.drawCentredString(w / 2, h - 4 * cm, f"Referencia: {referencia}")
    c.setFont("Helvetica-Bold", 18)
    # nombre en varias líneas si es largo
    y = h / 2 + 1 * cm
    palabras, linea, lineas = nombre_doc.split(), "", []
    for p in palabras:
        if len(linea + " " + p) > 42:
            lineas.append(linea.strip())
            linea = p
        else:
            linea += " " + p
    lineas.append(linea.strip())
    for ln in lineas[:5]:
        c.drawCentredString(w / 2, y, ln)
        y -= 0.9 * cm
    if pendiente:
        c.setFillColorRGB(0.75, 0.1, 0.1)
        c.setFont("Helvetica-Bold", 13)
        c.drawCentredString(w / 2, y - 1 * cm, "[[ PENDIENTE — DOCUMENTO NO ADJUNTO ]]")
        if motivo:
            c.setFont("Helvetica", 10)
            c.drawCentredString(w / 2, y - 1.8 * cm, motivo[:95])
        c.setFillColorRGB(0, 0, 0)
    c.setFont("Helvetica", 8)
    c.drawCentredString(w / 2, 2 * cm, "Generado por LicitacionLab — Expediente Digital")
    c.showPage()
    c.save()
    return buf.getvalue()


def caratula_e_indice(referencia: str, institucion: str | None,
                      empresa: str, indice: list[dict],
                      pendientes: list[dict],
                      paginas_por_pieza: dict | None = None) -> bytes:
    """
    Carátula + índice + hoja de pendientes en PDF.
    Si paginas_por_pieza viene (modo PDF merged), el índice sale paginado
    con números REALES contados después del merge.
    """
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf, pagesize=letter)
    w, h = letter

    # Carátula
    c.setFont("Helvetica-Bold", 22)
    c.drawCentredString(w / 2, h - 8 * cm, "SOBRE A — OFERTA TÉCNICA")
    c.setFont("Helvetica-Bold", 13)
    c.drawCentredString(w / 2, h - 10 * cm, (empresa or "").upper())
    c.setFont("Helvetica", 12)
    c.drawCentredString(w / 2, h - 11.2 * cm, (institucion or "").upper())
    c.drawCentredString(w / 2, h - 12.2 * cm, f"Referencia: {referencia}")
    c.drawCentredString(w / 2, h - 13.2 * cm, f"Fecha: {_hoy()}")
    c.showPage()

    # Índice
    c.setFont("Helvetica-Bold", 15)
    c.drawString(2.5 * cm, h - 2.5 * cm, "Índice del Sobre A")
    c.setFont("Helvetica", 10)
    y = h - 3.6 * cm
    for i, item in enumerate(indice, 1):
        if y < 2.5 * cm:
            c.showPage()
            c.setFont("Helvetica", 10)
            y = h - 2.5 * cm
        pag = ""
        if paginas_por_pieza and item["id"] in paginas_por_pieza:
            pag = f"  pág. {paginas_por_pieza[item['id']]}"
        c.drawString(2.5 * cm, y, f"{i:>2}. {item['nombre'][:80]}{pag}")
        y -= 0.55 * cm
    c.showPage()

    # Hoja de pendientes
    if pendientes:
        c.setFont("Helvetica-Bold", 15)
        c.setFillColorRGB(0.75, 0.1, 0.1)
        c.drawString(2.5 * cm, h - 2.5 * cm, "Requisitos pendientes")
        c.setFillColorRGB(0, 0, 0)
        c.setFont("Helvetica", 10)
        y = h - 3.6 * cm
        for p in pendientes:
            if y < 2.5 * cm:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = h - 2.5 * cm
            c.drawString(2.5 * cm, y, f"• {p['nombre'][:70]} — {p.get('motivo') or 'pendiente'}"[:110])
            y -= 0.55 * cm
        c.showPage()
    c.save()
    return buf.getvalue()


# ══════════════════════════════════════════════════════════════
# Conversión docx → PDF (opcional, requiere LibreOffice)
# ══════════════════════════════════════════════════════════════

def soffice_disponible() -> bool:
    return shutil.which("soffice") is not None or shutil.which("libreoffice") is not None


def docx_a_pdf(docx_bytes: bytes) -> bytes | None:
    """Convierte con LibreOffice headless; None si no está disponible o falla."""
    binario = shutil.which("soffice") or shutil.which("libreoffice")
    if not binario:
        return None
    with tempfile.TemporaryDirectory() as tmp:
        origen = os.path.join(tmp, "f.docx")
        with open(origen, "wb") as f:
            f.write(docx_bytes)
        try:
            subprocess.run(
                [binario, "--headless", "--convert-to", "pdf", "--outdir", tmp, origen],
                check=True, capture_output=True, timeout=120,
            )
            destino = os.path.join(tmp, "f.pdf")
            if os.path.exists(destino):
                with open(destino, "rb") as f:
                    return f.read()
        except (subprocess.SubprocessError, OSError):
            pass
    return None


# ══════════════════════════════════════════════════════════════
# Ensamblado del paquete
# ══════════════════════════════════════════════════════════════

def construir_paquete(sb, empresa_id: str, referencia: str,
                      piezas_solicitadas: list[str] | str,
                      formato: str = "docx",
                      progreso_cb=None) -> dict:
    """
    Genera el ZIP del Sobre A. Devuelve:
      {"zip_bytes", "nombre", "indice", "pendientes", "omitidas"}
    - piezas_solicitadas: lista de ids del catálogo o "todas".
    - formato 'pdf': formularios convertidos con LibreOffice si está;
      además intenta un merge completo Sobre_A_completo.pdf con índice
      paginado real. Si soffice no está, cae a docx sin romper el job.
    - progreso_cb(pieza_actual, completadas, total): hook del worker.
    - Regla D2: piezas 'err' por vencimiento/rechazo se OMITEN del ZIP y
      van a la hoja de pendientes; si el doc simplemente falta, entra
      SOLO su hoja separadora marcada.
    """
    ctx = datos_contexto(sb, empresa_id, referencia)
    catalogo = listar_piezas(sb, empresa_id, referencia)
    if piezas_solicitadas == "todas" or "todas" in piezas_solicitadas:
        elegidas = catalogo
    else:
        pedidos = set(piezas_solicitadas)
        elegidas = [p for p in catalogo if p["id"] in pedidos]

    usar_pdf = formato == "pdf" and soffice_disponible()
    total = len(elegidas)
    indice, pendientes, omitidas = [], [], []
    entradas: list[tuple[str, bytes, str]] = []  # (ruta_zip, bytes, pieza_id)

    for n, pieza in enumerate(elegidas, 1):
        if progreso_cb:
            progreso_cb(pieza["nombre"], n - 1, total)

        if pieza["tipo"] == "formulario":
            contenido = GENERADORES[pieza["id"]](sb, empresa_id, ctx)
            ext = "docx"
            if usar_pdf:
                pdf = docx_a_pdf(contenido)
                if pdf:
                    contenido, ext = pdf, "pdf"
            nombre_archivo = f"{n:02d}_{_slug(pieza['nombre'])}.{ext}"
            entradas.append((nombre_archivo, contenido, pieza["id"]))
            indice.append(pieza)
            if pieza["estado"] == "warn":
                pendientes.append(pieza)
            continue

        # --- documento del expediente ---
        if pieza["estado"] == "err" and "vencido" in (pieza.get("motivo") or "").lower():
            omitidas.append(pieza)          # JAMÁS entra al ZIP (D2)
            pendientes.append(pieza)
            continue

        sep = hoja_separadora(pieza["nombre"], referencia, ctx.get("institucion"),
                              pendiente=pieza["estado"] == "err",
                              motivo=pieza.get("motivo"))
        entradas.append((f"Documentos/{n:02d}_Separador_{_slug(pieza['nombre'])}.pdf",
                         sep, pieza["id"] + ":sep"))

        contenido_doc = None
        if pieza["estado"] != "err":
            contenido_doc = _bytes_de_documento(sb, empresa_id, pieza)
        if contenido_doc:
            nombre_archivo = f"Documentos/{n:02d}_{_slug(pieza['nombre'])}.pdf"
            entradas.append((nombre_archivo, contenido_doc, pieza["id"]))
            indice.append(pieza)
            if pieza["estado"] == "warn":
                pendientes.append(pieza)
        else:
            indice.append(pieza | {"nombre": pieza["nombre"] + " (solo separador)"})
            pendientes.append(pieza)

    if progreso_cb:
        progreso_cb("Carátula e índice", total, total)

    empresa_nombre = _v(ctx["empresa"], "razon_social", "nombre_perfil") or ""
    fecha = date.today().strftime("%Y%m%d")
    nombre_zip = f"Sobre_A_{_slug(referencia)}_{fecha}.zip"

    # Merge PDF completo con índice paginado real (solo modo pdf)
    merged, paginas_por_pieza = None, None
    if usar_pdf:
        merged, paginas_por_pieza = _merge_completo(entradas)

    caratula = caratula_e_indice(referencia, ctx.get("institucion"),
                                 empresa_nombre, indice, pendientes,
                                 paginas_por_pieza)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("00_Caratula_e_Indice.pdf", caratula)
        for ruta, contenido, _pid in entradas:
            z.writestr(ruta, contenido)
        if merged:
            # el merged final lleva la carátula (con paginado real) al frente
            final = PdfWriter()
            for pag in PdfReader(io.BytesIO(caratula)).pages:
                final.add_page(pag)
            for pag in PdfReader(io.BytesIO(merged)).pages:
                final.add_page(pag)
            out = io.BytesIO()
            final.write(out)
            z.writestr("Sobre_A_completo.pdf", out.getvalue())

    return {"zip_bytes": buf.getvalue(), "nombre": nombre_zip,
            "indice": [{"id": p["id"], "nombre": p["nombre"], "estado": p["estado"]}
                       for p in indice],
            "pendientes": [{"id": p["id"], "nombre": p["nombre"],
                            "motivo": p.get("motivo")} for p in pendientes],
            "omitidas": [{"id": p["id"], "nombre": p["nombre"],
                          "motivo": p.get("motivo")} for p in omitidas]}


def _bytes_de_documento(sb, empresa_id: str, pieza: dict) -> bytes | None:
    """Descarga el adjunto de una pieza 'documento' (bid_archivos o legacy)."""
    if pieza.get("_legacy_path"):
        return _descargar_binario(sb, pieza["_legacy_path"])
    archivo_id = pieza.get("archivo_id")
    if not archivo_id or str(archivo_id).startswith("legacy:"):
        return None
    reg = (sb.table("bid_archivos").select("storage_path,nombre")
           .eq("id", archivo_id).eq("empresa_id", empresa_id)
           .limit(1).execute().data or [None])[0]
    if not reg:
        return None
    return _descargar_binario(sb, reg.get("storage_path") or "")


def _merge_completo(entradas) -> tuple[bytes | None, dict | None]:
    """
    Une todos los PDFs en un solo documento y cuenta la página REAL de
    arranque de cada pieza (para el índice paginado — regla del handoff:
    paginar DESPUÉS del merge).
    """
    writer = PdfWriter()
    paginas: dict[str, int] = {}
    pagina_actual = 1
    for _ruta, contenido, pid in entradas:
        try:
            reader = PdfReader(io.BytesIO(contenido))
        except Exception:
            continue  # docx u otro no-PDF: no entra al merged
        if not pid.endswith(":sep"):
            paginas[pid] = pagina_actual
        for pag in reader.pages:
            writer.add_page(pag)
            pagina_actual += 1
    if pagina_actual == 1:
        return None, None
    out = io.BytesIO()
    writer.write(out)
    return out.getvalue(), paginas
