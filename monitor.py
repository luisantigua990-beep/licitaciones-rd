"""
Monitor de Procesos de Compras Públicas - República Dominicana
==============================================================
Módulo que consulta la API de la DGCP, detecta procesos nuevos,
y guarda procesos + artículos (con códigos UNSPSC) en Supabase.
"""

import os
import time
import requests
from datetime import datetime, timedelta
from dotenv import load_dotenv
from supabase import create_client

load_dotenv()

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
API_BASE_URL = "https://datosabiertos.dgcp.gob.do/api-dgcp/v1"

# Ciclo de refresco del Data Warehouse de la DGCP (cada 8 horas)
CICLO_HORAS_DGCP = 8

supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

# ============================================
# CRON LOG
# ============================================

def registrar_cron_log(job: str, status: str = "ok", detalle: dict = None, duracion_ms: int = None):
    """Registra ejecución del job en cron_log para auditoría."""
    try:
        supabase.table("cron_log").insert({
            "job": job,
            "status": status,
            "detalle": detalle or {},
            "duracion_ms": duracion_ms,
        }).execute()
    except Exception as e:
        print(f"⚠️  No se pudo registrar cron_log: {e}")



# ============================================
# TRACKING DE SINCRONIZACIÓN
# ============================================

def registrar_sync_exitosa(procesos_encontrados: int, procesos_nuevos: int):
    """Guarda el timestamp de la última sincronización exitosa con la API DGCP."""
    ahora = datetime.utcnow().isoformat()
    try:
        supabase.table("system_config").upsert({
            "clave": "ultima_sync_dgcp",
            "valor": ahora,
            "meta": {
                "procesos_encontrados": procesos_encontrados,
                "procesos_nuevos": procesos_nuevos,
                "ciclo_horas": CICLO_HORAS_DGCP
            }
        }, on_conflict="clave").execute()
        print(f"📡 Sync registrada: {ahora}")
    except Exception as e:
        print(f"⚠️  No se pudo registrar sync (tabla system_config puede no existir): {e}")


def obtener_estado_sync():
    """
    Devuelve el estado de sincronización con la API DGCP.
    Retorna: ultima_sync, proxima_sync, minutos_desde_sync, ciclo_horas
    """
    try:
        result = supabase.table("system_config") \
            .select("valor, meta") \
            .eq("clave", "ultima_sync_dgcp") \
            .single() \
            .execute()

        if not result.data:
            return None

        ultima_sync_str = result.data["valor"]
        meta = result.data.get("meta", {})

        ultima_sync = datetime.fromisoformat(ultima_sync_str)
        proxima_sync = ultima_sync + timedelta(hours=CICLO_HORAS_DGCP)
        ahora = datetime.utcnow()
        minutos_desde_sync = int((ahora - ultima_sync).total_seconds() / 60)

        return {
            "ultima_sync": ultima_sync_str,
            "proxima_sync": proxima_sync.isoformat(),
            "minutos_desde_sync": minutos_desde_sync,
            "ciclo_horas": CICLO_HORAS_DGCP,
            "procesos_encontrados_ultimo_ciclo": meta.get("procesos_encontrados", 0),
            "procesos_nuevos_ultimo_ciclo": meta.get("procesos_nuevos", 0),
            "mensaje": f"Datos verificados hace {minutos_desde_sync} min. Próximo ciclo DGCP: {proxima_sync.strftime('%H:%M')} UTC"
        }

    except Exception as e:
        print(f"⚠️  Error obteniendo estado sync: {e}")
        return None


# ============================================
# FUNCIONES PARA PROCESOS
# ============================================

def obtener_procesos_api(fecha_desde=None, fecha_hasta=None, page=1, limit=1000):
    if fecha_desde is None:
        fecha_desde = (datetime.now() - timedelta(days=1)).strftime("%Y-%m-%d")
    if fecha_hasta is None:
        fecha_hasta = datetime.now().strftime("%Y-%m-%d")

    params = {"startdate": fecha_desde, "enddate": fecha_hasta, "page": page, "limit": limit}

    try:
        response = requests.get(f"{API_BASE_URL}/procesos", params=params, timeout=30)
        response.raise_for_status()
        data = response.json()

        if data.get("hasError"):
            return []

        payload = data.get("payload")
        if payload is None:
            return []

        if isinstance(payload, list):
            procesos = payload
        elif isinstance(payload, dict):
            procesos = payload.get("content", []) or []
        else:
            procesos = []

        print(f"📥 API: {len(procesos)} procesos (pág {page})")
        return procesos

    except requests.exceptions.RequestException as e:
        print(f"❌ Error API: {e}")
        return []


def obtener_todas_las_paginas(fecha_desde=None, fecha_hasta=None):
    todos = []
    page = 1
    while True:
        procesos = obtener_procesos_api(fecha_desde, fecha_hasta, page=page, limit=1000)
        if not procesos:
            break
        todos.extend(procesos)
        if len(procesos) < 1000:
            break
        page += 1
    return todos


def obtener_codigos_existentes(codigos):
    if not codigos:
        return set()
    try:
        existentes = set()
        for i in range(0, len(codigos), 50):
            bloque = codigos[i:i+50]
            result = supabase.table("procesos").select("codigo_proceso").in_("codigo_proceso", bloque).execute()
            existentes.update(r["codigo_proceso"] for r in result.data)
        return existentes
    except Exception as e:
        print(f"❌ Error Supabase: {e}")
        return set()


def guardar_procesos_nuevos(procesos):
    """
    Para cada proceso de la API DGCP:
    - Si NO existe en Supabase → insertar como nuevo
    - Si YA existe (vino del scraper) → actualizar con los campos que faltan
    """
    if not procesos:
        return []

    codigos = [p["codigo_proceso"] for p in procesos]

    # Obtener cuáles ya existen y si están enriquecidos
    try:
        existentes_data = {}
        for i in range(0, len(codigos), 50):
            bloque = codigos[i:i+50]
            result = supabase.table("procesos") \
                .select("codigo_proceso, enriquecido_api") \
                .in_("codigo_proceso", bloque) \
                .execute()
            for r in result.data:
                existentes_data[r["codigo_proceso"]] = r.get("enriquecido_api", False)
    except Exception as e:
        print(f"❌ Error consultando existentes: {e}")
        return []

    nuevos = []
    para_insertar = []
    para_actualizar = []

    for p in procesos:
        codigo = p["codigo_proceso"]
        if codigo not in existentes_data:
            # Proceso nuevo — insertar completo
            para_insertar.append({
                "codigo_proceso":                p.get("codigo_proceso"),
                "codigo_unidad_compra":          p.get("codigo_unidad_compra"),
                "unidad_compra":                 p.get("unidad_compra"),
                "modalidad":                     p.get("modalidad"),
                "tipo_excepcion":                p.get("tipo_excepcion"),
                "titulo":                        p.get("titulo", "").strip() if p.get("titulo") else None,
                "descripcion":                   p.get("descripcion", "").strip() if p.get("descripcion") else None,
                "estado_proceso":                p.get("estado_proceso"),
                "divisa":                        p.get("divisa", "DOP"),
                "monto_estimado":                p.get("monto_estimado"),
                "fecha_publicacion":             p.get("fecha_publicacion"),
                "fecha_enmienda":                p.get("fecha_enmienda"),
                "fecha_fin_recepcion_ofertas":   p.get("fecha_fin_recepcion_ofertas"),
                "fecha_apertura_ofertas":        p.get("fecha_apertura_ofertas"),
                "fecha_estimada_adjudicacion":   p.get("fecha_estimada_adjudicacion"),
                "fecha_suscripcion":             p.get("fecha_suscripcion"),
                "dirigido_mipymes":              p.get("dirigido_mipymes"),
                "dirigido_mipymes_mujeres":      p.get("dirigido_mipymes_mujeres"),
                "objeto_proceso":                p.get("objeto_proceso"),
                "subobjeto_proceso":             p.get("subobjeto_proceso"),
                "url":                           p.get("url"),
                "duracion_contrato":             p.get("duracion_contrato"),
                "proceso_lotificado":            p.get("proceso_lotificado"),
                "compra_verde":                  p.get("compra_verde"),
                "compra_conjunta":               p.get("compra_conjunta"),
                "notificado":                    False,
                "enriquecido_api":               True,
            })
            nuevos.append(p)
        elif not existentes_data[codigo]:
            # Existe pero vino del scraper (enriquecido_api=False) → actualizar
            para_actualizar.append({
                "codigo_proceso":              codigo,
                "objeto_proceso":              p.get("objeto_proceso"),
                "subobjeto_proceso":           p.get("subobjeto_proceso"),
                "modalidad":                   p.get("modalidad"),
                "tipo_excepcion":              p.get("tipo_excepcion"),
                "estado_proceso":              p.get("estado_proceso"),
                "divisa":                      p.get("divisa", "DOP"),
                "fecha_enmienda":              p.get("fecha_enmienda"),
                "fecha_apertura_ofertas":      p.get("fecha_apertura_ofertas"),
                "fecha_estimada_adjudicacion": p.get("fecha_estimada_adjudicacion"),
                "fecha_suscripcion":           p.get("fecha_suscripcion"),
                "dirigido_mipymes":            p.get("dirigido_mipymes"),
                "dirigido_mipymes_mujeres":    p.get("dirigido_mipymes_mujeres"),
                "duracion_contrato":           p.get("duracion_contrato"),
                "proceso_lotificado":          p.get("proceso_lotificado"),
                "compra_verde":                p.get("compra_verde"),
                "compra_conjunta":             p.get("compra_conjunta"),
                "url":                         p.get("url"),
                "enriquecido_api":             True,
            })

    # Insertar nuevos
    if para_insertar:
        try:
            insertados = 0
            for i in range(0, len(para_insertar), 100):
                result = supabase.table("procesos").insert(para_insertar[i:i+100]).execute()
                insertados += len(result.data)
            print(f"✅ {insertados} procesos nuevos insertados")
        except Exception as e:
            print(f"❌ Error insertando procesos: {e}")

    # Actualizar los que vinieron del scraper
    if para_actualizar:
        try:
            actualizados = 0
            for reg in para_actualizar:
                codigo = reg.pop("codigo_proceso")
                supabase.table("procesos").update(reg).eq("codigo_proceso", codigo).execute()
                actualizados += 1
            print(f"🔄 {actualizados} procesos enriquecidos con datos de la API")
        except Exception as e:
            print(f"❌ Error actualizando procesos: {e}")

    if not nuevos:
        print("ℹ️  Sin procesos nuevos (todos ya existían o fueron enriquecidos)")

    return nuevos


# ============================================
# FUNCIONES PARA ARTÍCULOS
# ============================================

def obtener_articulos_proceso(codigo_proceso):
    try:
        response = requests.get(
            f"{API_BASE_URL}/procesos/articulos",
            params={"proceso": codigo_proceso, "limit": 1000},
            timeout=30
        )
        response.raise_for_status()
        data = response.json()
        payload = data.get("payload")
        if payload is None:
            return []
        if isinstance(payload, list):
            return payload
        elif isinstance(payload, dict):
            return payload.get("content", []) or []
        return []
    except Exception:
        return []


def guardar_articulos(codigo_proceso, articulos):
    if not articulos:
        return 0
    registros = [{
        "codigo_proceso": codigo_proceso,
        "familia_unspsc": a.get("familia_unspsc"),
        "clase_unspsc": a.get("clase_unspsc"),
        "subclase_unspsc": a.get("subclase_unspsc"),
        "descripcion_articulo": a.get("descripcion_articulo"),
        "descripcion_usuario": a.get("descripcion_usuario"),
        "cuenta_presupuestaria": a.get("cuenta_presupuestaria"),
        "cantidad": a.get("cantidad"),
        "unidad_medida": a.get("unidad_medida"),
        "precio_unitario_estimado": a.get("precio_unitario_estimado"),
        "precio_total_estimado": a.get("precio_total_estimado"),
    } for a in articulos]

    try:
        insertados = 0
        for i in range(0, len(registros), 100):
            bloque = registros[i:i+100]
            result = supabase.table("articulos_proceso").insert(bloque).execute()
            insertados += len(result.data)
        return insertados
    except Exception:
        return 0


def procesar_articulos_de_nuevos(procesos_nuevos):
    """Obtiene y guarda artículos UNSPSC para procesos nuevos."""
    if not procesos_nuevos:
        return
    total_articulos = 0
    for proceso in procesos_nuevos:
        articulos = obtener_articulos_proceso(proceso["codigo_proceso"])
        if articulos:
            total_articulos += guardar_articulos(proceso["codigo_proceso"], articulos)
        time.sleep(0.2)
    print(f"✅ {total_articulos} artículos guardados")


def enriquecer_articulos_faltantes():
    """
    VÍA 2 — Fallback de reconciliación (corre cada 8h junto al ETL de la DGCP).

    Busca procesos del scraper que quedaron incompletos (sin artículos o sin noticeUID)
    y los completa usando la API de Datos Abiertos, que a estas alturas ya fue actualizada.

    También reconcilia procesos enriquecidos por la API que aún no tienen artículos.
    """
    try:
        from datetime import timedelta

        # ── A. Reconciliar procesos del scraper sin artículos (enriquecido_api=False) ──
        limite_fecha = (datetime.now() - timedelta(hours=9)).isoformat()  # publicados hace >9h
        scraper_sin_arts = supabase.table("procesos")             .select("codigo_proceso, url, titulo")             .eq("enriquecido_api", False)             .lt("fecha_publicacion", limite_fecha)             .execute().data or []

        if scraper_sin_arts:
            codigos_scraper = [p["codigo_proceso"] for p in scraper_sin_arts]
            # Filtrar los que realmente no tienen artículos
            con_arts = set()
            for i in range(0, len(codigos_scraper), 50):
                r = supabase.table("articulos_proceso").select("codigo_proceso")                     .in_("codigo_proceso", codigos_scraper[i:i+50]).execute()
                con_arts.update(x["codigo_proceso"] for x in r.data)

            incompletos = [p for p in scraper_sin_arts if p["codigo_proceso"] not in con_arts]

            if incompletos:
                print(f"🔄 VÍA 2: {len(incompletos)} procesos del scraper incompletos — reconciliando con API...")
                total_rec = 0
                for proc in incompletos[:50]:
                    codigo = proc["codigo_proceso"]

                    # Intentar obtener artículos de la API (ETL ya corrió)
                    articulos = obtener_articulos_proceso(codigo)
                    if articulos:
                        guardados = guardar_articulos(codigo, articulos)
                        total_rec += guardados

                    # También reconciliar la URL con noticeUID si aún no la tiene
                    url_actual = proc.get("url", "")
                    if "noticeUID" not in url_actual:
                        # Intentar obtener la URL con noticeUID desde la API de documentos
                        try:
                            resp = requests.get(
                                f"{API_BASE_URL}/procesos/documentos",
                                params={"proceso": codigo},
                                timeout=10
                            )
                            docs_data = resp.json()
                            docs = docs_data.get("payload", {})
                            if isinstance(docs, dict):
                                docs = docs.get("content", [])
                            elif not isinstance(docs, list):
                                docs = []

                            for doc in docs:
                                url_doc = doc.get("url_documento", "")
                                if url_doc and "noticeUID" in url_doc:
                                    # Extraer noticeUID de la URL del documento
                                    import re as _re
                                    m = _re.search(r"noticeUID=(DO1\.NTC\.[\w\.]+)", url_doc)
                                    if m:
                                        notice_uid = m.group(1)
                                        url_detalle = (
                                            f"https://comunidad.comprasdominicana.gob.do"
                                            f"/Public/Tendering/OpportunityDetail/Index"
                                            f"?noticeUID={notice_uid}"
                                        )
                                        supabase.table("procesos").update({"url": url_detalle})                                             .eq("codigo_proceso", codigo).execute()
                                        print(f"   🔗 [{codigo}] URL reconciliada: noticeUID={notice_uid}")
                                        break
                        except Exception:
                            pass

                    # Marcar como enriquecido si ya tiene artículos
                    if articulos:
                        supabase.table("procesos").update({"enriquecido_api": True})                             .eq("codigo_proceso", codigo).execute()

                    time.sleep(0.3)

                print(f"✅ VÍA 2: {total_rec} artículos reconciliados para {len(incompletos)} procesos")

        # ── B. Procesos ya marcados enriquecido_api=True pero sin artículos ──
        enriquecidos = supabase.table("procesos")             .select("codigo_proceso")             .eq("enriquecido_api", True)             .execute().data or []

        if not enriquecidos:
            return

        codigos_enriquecidos = [r["codigo_proceso"] for r in enriquecidos]
        con_articulos = set()
        for i in range(0, len(codigos_enriquecidos), 50):
            r = supabase.table("articulos_proceso").select("codigo_proceso")                 .in_("codigo_proceso", codigos_enriquecidos[i:i+50]).execute()
            con_articulos.update(x["codigo_proceso"] for x in r.data)

        sin_articulos = [c for c in codigos_enriquecidos if c not in con_articulos]

        if not sin_articulos:
            return

        print(f"🔎 {len(sin_articulos)} procesos enriquecidos sin artículos — completando...")
        total = 0
        for codigo in sin_articulos[:50]:
            articulos = obtener_articulos_proceso(codigo)
            if articulos:
                total += guardar_articulos(codigo, articulos)
            time.sleep(0.2)

        if total:
            print(f"✅ {total} artículos completados en procesos existentes")

    except Exception as e:
        print(f"⚠️  Error en enriquecer_articulos_faltantes: {e}")


# ============================================
# NOTIFICACIONES
# ============================================

def notificar_procesos_nuevos(procesos_nuevos):
    if not procesos_nuevos:
        return

    try:
        from notifications import enviar_notificacion
        from pywebpush import WebPushException

        result = supabase.table("user_subscriptions")            .select("*")            .eq("active", True)            .execute()

        suscripciones = result.data
        if not suscripciones:
            print("ℹ️  Sin suscriptores activos para notificar")
            return

        print(f"📋 {len(suscripciones)} suscriptores activos, {len(procesos_nuevos)} proceso(s) nuevo(s)")

        APP_URL = os.getenv("APP_URL", "https://app.licitacionlab.com")
        notificaciones_enviadas = 0
        errores = 0

        for proceso in procesos_nuevos:
            titulo_proceso = proceso.get("titulo", "Nueva licitación")[:70]
            entidad        = (proceso.get("unidad_compra") or "")[:40]
            monto          = proceso.get("monto_estimado")
            monto_str      = f" | RD${monto:,.0f}" if monto else ""
            codigo         = proceso.get("codigo_proceso", "")
            url_proceso    = f"{APP_URL}?proceso={codigo}"

            for sub in suscripciones:
                subscription_info = {
                    "endpoint": sub["endpoint"],
                    "keys": {"auth": sub["auth"], "p256dh": sub["p256dh"]}
                }
                try:
                    resultado = enviar_notificacion(
                        subscription_info,
                        titulo=f"🏗️ {entidad}{monto_str}",
                        cuerpo=titulo_proceso,
                        url=url_proceso
                    )
                    if resultado is True:
                        notificaciones_enviadas += 1
                    elif resultado == "410":
                        # Suscripción expirada — desactivar
                        supabase.table("user_subscriptions")                            .update({"active": False})                            .eq("endpoint", sub["endpoint"]).execute()
                        print(f"🧹 Suscripción 410 desactivada: {sub['endpoint'][:50]}...")
                    else:
                        errores += 1
                        print(f"⚠️  Push falló (False) para {sub['endpoint'][:50]}...")
                except WebPushException as e:
                    errores += 1
                    resp = getattr(e, "response", None)
                    status = getattr(resp, "status_code", "?") if resp else "?"
                    print(f"❌ WebPushException [{status}]: {str(e)[:120]}")
                except Exception as e:
                    errores += 1
                    print(f"❌ Error push inesperado: {type(e).__name__}: {str(e)[:120]}")

        print(f"🔔 Push: {notificaciones_enviadas} enviadas, {errores} errores")

    except Exception as e:
        import traceback
        print(f"❌ Error en notificar_procesos_nuevos: {e}")
        traceback.print_exc()


# ============================================
# FUNCIÓN PRINCIPAL
# ============================================


# ============================================
# MONITOR DE ENMIENDAS
# ============================================

GEMINI_API_KEY  = os.getenv("GEMINI_API_KEY", "")
RESEND_API_KEY  = os.getenv("RESEND_API_KEY", "")
SUPABASE_URL_ENV = os.getenv("SUPABASE_URL", "")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY", os.getenv("SUPABASE_KEY", ""))
APP_URL         = os.getenv("APP_URL", "https://web-production-7b940.up.railway.app")
FROM_EMAIL      = os.getenv("RESEND_FROM", "LicitacionLab <notificaciones@licitacionlab.com>")
TEST_EMAIL      = os.getenv("RESEND_TEST_EMAIL", "")


def _obtener_documentos_portal(noticeUID: str) -> list[dict]:
    """
    Scrape la página de documentos del proceso en el portal y devuelve
    lista de {nombre, document_file_id} de todos los archivos visibles.
    """
    import re
    from bs4 import BeautifulSoup

    url = (
        f"https://comunidad.comprasdominicana.gob.do/Public/Tendering/"
        f"OpportunityDetail/Index?noticeUID={noticeUID}&isModal=true&asPopupView=true"
    )
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "text/html,application/xhtml+xml,*/*;q=0.9",
        "Accept-Language": "es-DO,es;q=0.9",
    }
    try:
        resp = requests.get(url, headers=headers, timeout=30)
        if resp.status_code != 200:
            return []
        soup = BeautifulSoup(resp.text, "html.parser")
        documentos = []
        # Buscar links de descarga con documentFileId
        for a in soup.find_all("a", href=True):
            href = a["href"]
            m = re.search(r"documentFileId=(\d+)", href)
            if m:
                nombre = a.get_text(strip=True) or f"documento_{m.group(1)}"
                documentos.append({
                    "nombre": nombre,
                    "document_file_id": m.group(1),
                    "url": href if href.startswith("http") else
                           f"https://comunidad.comprasdominicana.gob.do{href}"
                })
        return documentos
    except Exception as e:
        print(f"   ⚠️ Error scraping documentos {noticeUID}: {e}")
        return []


def _descargar_texto_enmienda(url_doc: str, nombre: str) -> str:
    """Descarga un documento (PDF o DOCX) y extrae su texto."""
    import io
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        PdfReader = None
    try:
        from docx import Document as DocxDocument
        DOCX_OK = True
    except ImportError:
        DOCX_OK = False

    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36",
        "Accept": "*/*",
    }
    try:
        # El portal usa redirect JS igual que en el análisis principal
        session = requests.Session()
        resp = session.get(url_doc, headers=headers, timeout=60)
        if len(resp.content) < 2000 and b"window.location.href" in resp.content:
            import re
            m = re.search('window.location.href\\s*=\\s*["\'](.*?)["\']', resp.text)
            if m:
                url_real = m.group(1)
                if not url_real.startswith("http"):
                    url_real = f"https://comunidad.comprasdominicana.gob.do{url_real}"
                resp = session.get(url_real, headers=headers, timeout=60)

        if resp.status_code != 200 or len(resp.content) < 200:
            return ""

        doc_bytes = resp.content
        es_docx = (
            doc_bytes[:4] == b"PK\x03\x04"
            or nombre.lower().endswith(".docx")
            or nombre.lower().endswith(".doc")
        )

        if es_docx and DOCX_OK:
            doc = DocxDocument(io.BytesIO(doc_bytes))
            parrafos = [p.text for p in doc.paragraphs if p.text.strip()]
            for tabla in doc.tables:
                for fila in tabla.rows:
                    ft = " | ".join(c.text.strip() for c in fila.cells if c.text.strip())
                    if ft:
                        parrafos.append(ft)
            return "\n".join(parrafos)[:8000]

        if PdfReader:
            lector = PdfReader(io.BytesIO(doc_bytes))
            texto = ""
            for i in range(min(len(lector.pages), 20)):
                t = lector.pages[i].extract_text()
                if t:
                    texto += t + "\n"
            return texto[:8000]

    except Exception as e:
        print(f"   ⚠️ Error descargando enmienda {nombre}: {e}")
    return ""


def _analizar_enmienda_gemini(proceso_titulo: str, texto_enmienda: str, nombre_doc: str) -> str:
    """Usa Gemini para resumir qué cambió en la enmienda."""
    if not GEMINI_API_KEY or not texto_enmienda.strip():
        return "No se pudo analizar el contenido de la enmienda."
    try:
        from google import genai
        cliente = genai.Client(api_key=GEMINI_API_KEY)
        prompt = f"""Eres un experto en licitaciones públicas dominicanas.
Se publicó una ENMIENDA/ADENDA al proceso: "{proceso_titulo}"
Documento: {nombre_doc}

Contenido de la enmienda:
{texto_enmienda[:6000]}

Analiza y responde en español con este formato exacto:

RESUMEN: [Una oración resumiendo el cambio principal]

CAMBIOS CLAVE:
- [cambio 1]
- [cambio 2]
- [cambio 3 si aplica]

IMPACTO: [Alto/Medio/Bajo] — [Explicación breve de cómo afecta a los ofertantes]

ACCIÓN REQUERIDA: [Qué debe hacer el ofertante con esta información]
"""
        resp = cliente.models.generate_content(
            model="gemini-2.0-flash",
            contents=prompt
        )
        return resp.text.strip()
    except Exception as e:
        return f"Error en análisis Gemini: {e}"


def _obtener_emails_seguidores(proceso_codigo: str) -> list[tuple[str, str]]:
    """Retorna lista de (email, nombre) de usuarios que siguen el proceso en estado activo."""
    estados_activos = ["analizando", "preparando", "presentado"]
    try:
        seg = supabase.table("seguimiento_procesos")             .select("user_id")             .eq("proceso_codigo", proceso_codigo)             .in_("estado", estados_activos)             .execute()
        if not seg.data:
            return []

        user_ids = list(set(s["user_id"] for s in seg.data))
        resultado = []
        for uid in user_ids:
            try:
                resp_auth = requests.get(
                    f"{SUPABASE_URL_ENV}/auth/v1/admin/users/{uid}",
                    headers={
                        "apikey": SUPABASE_SERVICE_KEY,
                        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
                    },
                    timeout=10,
                )
                if resp_auth.status_code == 200:
                    user_data = resp_auth.json()
                    email = user_data.get("email")
                    meta  = user_data.get("user_metadata") or {}
                    nombre = (meta.get("nombre") or meta.get("full_name") or
                              meta.get("name") or (email.split("@")[0] if email else "Usuario"))
                    if email:
                        resultado.append((email, nombre))
            except Exception as e:
                print(f"   ⚠️ No se pudo obtener email de user {uid}: {e}")

        if TEST_EMAIL:
            return [(TEST_EMAIL, "Usuario de prueba")]
        return resultado
    except Exception as e:
        print(f"   ⚠️ Error obteniendo seguidores de {proceso_codigo}: {e}")
        return []


def _enviar_notificacion_push_enmienda(proceso_codigo: str, proceso_titulo: str, url_proceso: str):
    """Envía push notification a usuarios que siguen el proceso."""
    try:
        from notifications import enviar_notificacion
        seg = supabase.table("seguimiento_procesos")             .select("user_id")             .eq("proceso_codigo", proceso_codigo)             .in_("estado", ["analizando", "preparando", "presentado"])             .execute()
        if not seg.data:
            return
        user_ids = list(set(s["user_id"] for s in seg.data))

        subs = supabase.table("user_subscriptions")             .select("*")             .in_("user_id", user_ids)             .eq("active", True)             .execute()

        for sub in (subs.data or []):
            subscription_info = {
                "endpoint": sub["endpoint"],
                "keys": {"auth": sub["auth"], "p256dh": sub["p256dh"]}
            }
            enviar_notificacion(
                subscription_info,
                titulo=f"⚠️ Enmienda publicada — {proceso_codigo}",
                cuerpo=proceso_titulo[:80],
                url=url_proceso
            )
        print(f"   🔔 Push enviado a {len(subs.data or [])} suscriptores")
    except Exception as e:
        print(f"   ⚠️ Error enviando push enmienda: {e}")


def _enviar_email_enmienda(proceso_codigo: str, proceso_titulo: str,
                            nombre_doc: str, resumen_gemini: str,
                            url_proceso: str, emails: list[tuple[str, str]]):
    """Envía email con el resumen de la enmienda analizado por Gemini."""
    if not RESEND_API_KEY or not emails:
        return

    # Convertir texto Gemini a HTML legible
    html_resumen = ""
    for linea in resumen_gemini.split("\n"):
        linea = linea.strip()
        if not linea:
            continue
        if linea.startswith("RESUMEN:"):
            html_resumen += f"<p style=\'font-size:15px;color:#1e3a5f;font-weight:bold;margin:0 0 12px;\'>{linea}</p>"
        elif linea.startswith("CAMBIOS CLAVE:"):
            html_resumen += "<p style=\'font-size:13px;font-weight:bold;color:#374151;margin:16px 0 6px;\'>📋 Cambios clave:</p><ul style=\'margin:0;padding-left:20px;\'>"
        elif linea.startswith("IMPACTO:"):
            html_resumen += f"</ul><p style=\'font-size:13px;color:#374151;margin:12px 0;\'>⚡ <strong>{linea}</strong></p>"
        elif linea.startswith("ACCIÓN REQUERIDA:"):
            html_resumen += f"<p style=\'font-size:13px;color:#dc2626;font-weight:bold;margin:12px 0;\'>🎯 {linea}</p>"
        elif linea.startswith("- "):
            html_resumen += f"<li style=\'font-size:13px;color:#374151;margin:4px 0;\'>{linea[2:]}</li>"
        else:
            html_resumen += f"<p style=\'font-size:13px;color:#374151;margin:6px 0;\'>{linea}</p>"

    html_body = f"""<!DOCTYPE html>
<html><head><meta charset="UTF-8"></head>
<body style="margin:0;padding:0;background:#f3f4f6;font-family:\'Plus Jakarta Sans\',Arial,sans-serif;">
<div style="max-width:600px;margin:24px auto;background:white;border-radius:12px;overflow:hidden;box-shadow:0 4px 24px rgba(0,0,0,0.08);">
  <div style="background:linear-gradient(135deg,#d97706,#b45309);padding:28px 32px;">
    <p style="margin:0 0 4px;font-size:12px;color:rgba(255,255,255,0.8);letter-spacing:1px;">ENMIENDA DETECTADA</p>
    <h1 style="margin:0;font-size:20px;color:white;">⚠️ {proceso_codigo}</h1>
    <p style="margin:8px 0 0;font-size:14px;color:rgba(255,255,255,0.9);">{proceso_titulo[:90]}</p>
  </div>
  <div style="padding:24px 32px;">
    <div style="background:#fffbeb;border:1px solid #fbbf24;border-radius:8px;padding:14px 18px;margin-bottom:20px;">
      <p style="margin:0;font-size:13px;color:#92400e;">📄 Documento: <strong>{nombre_doc}</strong></p>
    </div>
    <h3 style="margin:0 0 14px;font-size:15px;color:#1e3a5f;">🤖 Análisis de la enmienda</h3>
    <div style="background:#f8fafc;border-radius:8px;padding:16px 20px;margin-bottom:20px;">
      {html_resumen}
    </div>
    <div style="text-align:center;margin-top:24px;">
      <a href="{url_proceso}" style="background:#1e3a5f;color:white;padding:12px 28px;border-radius:8px;text-decoration:none;font-size:14px;font-weight:bold;">Ver proceso completo →</a>
    </div>
  </div>
  <div style="background:#f8fafc;padding:16px 32px;border-top:1px solid #e5e7eb;">
    <p style="margin:0;font-size:11px;color:#9ca3af;text-align:center;">LicitacionLab — Monitoreo automático de enmiendas</p>
  </div>
</div>
</body></html>"""

    for email, nombre in emails:
        payload = {
            "from": FROM_EMAIL,
            "to": [email],
            "subject": f"⚠️ Enmienda en {proceso_codigo} — acción requerida",
            "html": html_body.replace("{{nombre}}", nombre),
        }
        try:
            resp = requests.post(
                "https://api.resend.com/emails",
                headers={
                    "Authorization": f"Bearer {RESEND_API_KEY}",
                    "Content-Type": "application/json",
                },
                json=payload,
                timeout=15,
            )
            if resp.status_code in (200, 201):
                print(f"   📧 Email enmienda enviado a {email}")
            else:
                print(f"   ⚠️ Error email enmienda a {email}: {resp.status_code}")
        except Exception as e:
            print(f"   ⚠️ Error enviando email enmienda: {e}")


def monitorear_enmiendas():
    """
    Verifica si los procesos en estado 'preparando' o posterior tienen
    documentos nuevos (enmiendas/adendas) desde la última verificación.
    Corre 2 veces al día desde ejecutar_monitor().
    """
    import time as _time
    t0 = _time.time()
    print(f"\n📋 Monitor de enmiendas | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    estados_activos = ["preparando", "presentado", "adjudicado"]

    try:
        # 1. Obtener procesos que alguien está siguiendo activamente
        seg = supabase.table("seguimiento_procesos")             .select("proceso_codigo")             .in_("estado", estados_activos)             .execute()

        if not seg.data:
            print("   ℹ️ Sin procesos activos en seguimiento")
            registrar_cron_log("monitor_enmiendas", "ok",
                {"nota": "sin procesos activos"}, int((_time.time()-t0)*1000))
            return

        codigos = list(set(s["proceso_codigo"] for s in seg.data))
        print(f"   📂 {len(codigos)} procesos en seguimiento activo")

        # 2. Obtener datos de esos procesos (necesitamos noticeUID y título)
        procesos_data = supabase.table("procesos")             .select("codigo_proceso, titulo, url, unidad_compra")             .in_("codigo_proceso", codigos)             .execute().data or []

        enmiendas_detectadas = 0

        for proc in procesos_data:
            codigo    = proc["codigo_proceso"]
            titulo    = proc.get("titulo", codigo)
            url_proc  = proc.get("url", "")

            # Extraer noticeUID de la URL
            import re
            m = re.search(r"noticeUID=(DO1\.NTC\.[\w\.]+)", url_proc or "")
            if not m:
                print(f"   ⚠️ [{codigo}] Sin noticeUID — omitiendo")
                continue
            notice_uid = m.group(1)

            # 3. Obtener documentos actuales del portal
            docs_actuales = _obtener_documentos_portal(notice_uid)
            if not docs_actuales:
                continue

            ids_actuales = {d["document_file_id"] for d in docs_actuales}

            # 4. Comparar con los IDs ya conocidos en Supabase
            registro = supabase.table("enmiendas_proceso")                 .select("document_file_ids_conocidos")                 .eq("codigo_proceso", codigo)                 .execute()

            if not registro.data:
                # Primera vez — guardar los IDs actuales como baseline, sin notificar
                supabase.table("enmiendas_proceso").insert({
                    "codigo_proceso": codigo,
                    "document_file_ids_conocidos": list(ids_actuales),
                    "ultima_verificacion": datetime.utcnow().isoformat(),
                }).execute()
                print(f"   📌 [{codigo}] Baseline creado ({len(ids_actuales)} docs)")
                continue

            ids_conocidos = set(registro.data[0]["document_file_ids_conocidos"])
            ids_nuevos    = ids_actuales - ids_conocidos

            if not ids_nuevos:
                # Actualizar timestamp de verificación
                supabase.table("enmiendas_proceso")                     .update({"ultima_verificacion": datetime.utcnow().isoformat()})                     .eq("codigo_proceso", codigo)                     .execute()
                continue

            # 5. Hay documentos nuevos — son enmiendas
            docs_nuevos = [d for d in docs_actuales if d["document_file_id"] in ids_nuevos]
            print(f"   🚨 [{codigo}] {len(docs_nuevos)} enmienda(s) detectada(s):")

            emails_seguidores = _obtener_emails_seguidores(codigo)
            url_proceso = url_proc or f"{APP_URL}/#proceso-{codigo}"

            for doc in docs_nuevos:
                nombre_doc = doc["nombre"]
                url_doc    = doc["url"]
                print(f"      📄 {nombre_doc}")

                # 6. Descargar y analizar con Gemini
                texto = _descargar_texto_enmienda(url_doc, nombre_doc)
                resumen_gemini = _analizar_enmienda_gemini(titulo, texto, nombre_doc)

                # 7. Guardar enmienda en Supabase
                supabase.table("enmiendas_proceso").upsert({
                    "codigo_proceso": codigo,
                    "document_file_ids_conocidos": list(ids_actuales),
                    "ultima_verificacion": datetime.utcnow().isoformat(),
                }).execute()

                supabase.table("enmiendas_detectadas").insert({
                    "codigo_proceso": codigo,
                    "nombre_documento": nombre_doc,
                    "document_file_id": doc["document_file_id"],
                    "url_documento": url_doc,
                    "resumen_gemini": resumen_gemini,
                    "fecha_deteccion": datetime.utcnow().isoformat(),
                }).execute()

                # 8. Push notification
                _enviar_notificacion_push_enmienda(codigo, titulo, url_proceso)

                # 9. Email con análisis
                _enviar_email_enmienda(codigo, titulo, nombre_doc,
                                       resumen_gemini, url_proceso, emails_seguidores)

                enmiendas_detectadas += 1

        registrar_cron_log("monitor_enmiendas", "ok", {
            "procesos_verificados": len(procesos_data),
            "enmiendas_detectadas": enmiendas_detectadas,
        }, int((_time.time()-t0)*1000))
        print(f"   ✅ Monitor enmiendas completado — {enmiendas_detectadas} enmienda(s) procesada(s)\n")

    except Exception as e:
        registrar_cron_log("monitor_enmiendas", "error", {"error": str(e)},
                           int((_time.time()-t0)*1000))
        print(f"   ❌ Error en monitor de enmiendas: {e}")


def ejecutar_monitor():
    import time as _time
    t0 = _time.time()
    print(f"\n{'='*50}")
    print(f"🔍 Monitor | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"{'='*50}")

    try:
        procesos = obtener_todas_las_paginas()
        if not procesos:
            print("📭 Sin procesos en el período")
            registrar_sync_exitosa(procesos_encontrados=0, procesos_nuevos=0)
            registrar_cron_log("monitor_api", "ok",
                {"procesos_encontrados": 0, "procesos_nuevos": 0, "nota": "sin procesos en período"},
                int((_time.time()-t0)*1000))
            return []

        print(f"📊 {len(procesos)} procesos encontrados")
        nuevos = guardar_procesos_nuevos(procesos)

        if nuevos:
            procesar_articulos_de_nuevos(nuevos)
            notificar_procesos_nuevos(nuevos)
            for p in nuevos[:5]:
                monto = f"RD${p.get('monto_estimado', 0):,.2f}" if p.get('monto_estimado') else "N/A"
                print(f"   🆕 {p['titulo'][:50]}... | {monto}")
            if len(nuevos) > 5:
                print(f"   ... y {len(nuevos) - 5} más")

        enriquecer_articulos_faltantes()

        # ── Monitor de enmiendas 2x/día (07-09h y 18-20h) ──
        hora_actual = datetime.now().hour
        if (7 <= hora_actual <= 9) or (18 <= hora_actual <= 20):
            monitorear_enmiendas()

        registrar_sync_exitosa(
            procesos_encontrados=len(procesos),
            procesos_nuevos=len(nuevos)
        )
        registrar_cron_log("monitor_api", "ok", {
            "procesos_encontrados": len(procesos),
            "procesos_nuevos": len(nuevos),
        }, int((_time.time()-t0)*1000))

        print(f"✅ Monitor completado\n")
        return nuevos

    except Exception as e:
        registrar_cron_log("monitor_api", "error", {"error": str(e)}, int((_time.time()-t0)*1000))
        raise


if __name__ == "__main__":
    ejecutar_monitor()
